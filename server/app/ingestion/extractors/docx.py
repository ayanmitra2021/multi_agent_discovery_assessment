import io
from typing import Optional

from docx import Document

from ...api.errors import DocumentParseError
from ...schemas.discovery import ParsedDocument, TableBlock

_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def extract_docx(
    content: bytes,
    filename: str,
    max_rows: Optional[int] = None,
) -> ParsedDocument:
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse DOCX '{filename}': {exc}") from exc

    text_blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    table_blocks: list[TableBlock] = []
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows[1:]
        ]
        warnings: list[str] = []
        if max_rows and len(rows) > max_rows:
            rows = rows[:max_rows]
            warnings.append(
                f"Table truncated to {max_rows} rows (MAX_APPS_PER_RUN limit)"
            )
        table_blocks.append(TableBlock(headers=headers, rows=rows))

    return ParsedDocument(
        filename=filename,
        content_type=_DOCX_CONTENT_TYPE,
        text_blocks=text_blocks,
        table_blocks=table_blocks,
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
        warnings=[],
    )
