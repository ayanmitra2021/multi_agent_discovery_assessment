"""In-memory document generators for tests. No files on disk needed."""
import csv
import io

import openpyxl
from docx import Document


def make_docx(
    paragraphs: list[str],
    tables: list[list[list[str]]] | None = None,
) -> bytes:
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    for table_data in tables or []:
        if not table_data:
            continue
        t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row in enumerate(table_data):
            for j, text in enumerate(row):
                t.cell(i, j).text = text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx(sheets: dict[str, list[list[str]]]) -> bytes:
    wb = openpyxl.Workbook()
    first = True
    for sheet_name, rows in sheets.items():
        if first:
            ws = wb.active
            ws.title = sheet_name
            first = False
        else:
            ws = wb.create_sheet(sheet_name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_csv(rows: list[list[str]]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    for row in rows:
        writer.writerow(row)
    return buf.getvalue().encode()
